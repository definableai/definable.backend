import asyncio
from pathlib import Path
from typing import IO, Any, List, Union
from uuid import uuid4

from libs.definable.document.base import Document
from libs.definable.document.reader.base import Reader

try:
  from docx import Document as DocxDocument  # type: ignore
except ImportError:
  raise ImportError("The `python-docx` package is not installed. Please install it via `pip install python-docx`.")


class DocxReader(Reader):
  """Reader for Doc/Docx files"""

  def read(self, file: Union[Path, IO[Any]]) -> List[Document]:
    """Read a docx file and return a list of documents"""
    try:
      if isinstance(file, Path):
        if not file.exists():
          raise FileNotFoundError(f"Could not find file: {file}")
        docx_document = DocxDocument(str(file))
        doc_name = file.stem
      else:
        docx_document = DocxDocument(file)
        doc_name = file.name.split(".")[0]

      doc_content = "\n\n".join([para.text for para in docx_document.paragraphs])

      documents = [
        Document(
          name=doc_name,
          id=str(uuid4()),
          content=doc_content,
        )
      ]

      if self.chunk:
        chunked_documents = []
        for document in documents:
          chunked_documents.extend(self.chunk_document(document))
        return chunked_documents
      return documents

    except Exception:
      return []

  async def async_read(self, file: Union[Path, IO[Any]]) -> List[Document]:
    """Asynchronously read a docx file and return a list of documents"""
    try:
      return await asyncio.to_thread(self.read, file)
    except Exception:
      return []
